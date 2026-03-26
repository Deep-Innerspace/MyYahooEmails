"""
Word document renderer using python-docx.

Renders a Report structure to a .docx file.
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from src.reports.builder import Report, ReportSection


# ── Severity colors ──────────────────────────────────────────────────────
_SEV_COLORS = {
    "high": RGBColor(0xDC, 0x26, 0x26),    # red
    "haute": RGBColor(0xDC, 0x26, 0x26),
    "medium": RGBColor(0xF5, 0x9E, 0x0B),  # orange
    "moyenne": RGBColor(0xF5, 0x9E, 0x0B),
    "low": RGBColor(0x6B, 0x72, 0x80),     # gray
    "basse": RGBColor(0x6B, 0x72, 0x80),
}

_HEADING_COLOR = RGBColor(0x1E, 0x3A, 0x5F)  # dark blue


def _add_heading(doc: Document, text: str, level: int):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=min(level, 4))
    for run in heading.runs:
        run.font.color.rgb = _HEADING_COLOR


def _add_table(doc: Document, headers: list, rows: list):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    # Color severity cells
                    text_lower = str(cell_text).lower().strip()
                    if text_lower in _SEV_COLORS:
                        run.font.color.rgb = _SEV_COLORS[text_lower]
                        run.font.bold = True


def _render_section(doc: Document, section: ReportSection):
    """Recursively render a section and its subsections."""
    _add_heading(doc, section.title, section.level)

    for para_text in section.paragraphs:
        p = doc.add_paragraph(para_text)
        p.style.font.size = Pt(10)

    if section.chart_path and section.chart_path.exists():
        doc.add_picture(str(section.chart_path), width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if section.table:
        _add_table(doc, section.table["headers"], section.table["rows"])

    for sub in section.subsections:
        _render_section(doc, sub)


def render_docx(report: Report, output_path: Path) -> Path:
    """Render a Report structure to a .docx file."""
    doc = Document()

    # ── Page setup ───────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title page ───────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(report.title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = _HEADING_COLOR

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run(report.subtitle)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Généré le {report.date}")
    date_run.font.size = Pt(10)
    date_run.font.italic = True

    doc.add_page_break()

    # ── Table of contents placeholder ────────────────────────────────────
    _add_heading(doc, "Table des matières", 1)
    for section in report.sections:
        indent = "  " * (section.level - 1)
        doc.add_paragraph(f"{indent}{section.title}")
    doc.add_page_break()

    # ── Sections ─────────────────────────────────────────────────────────
    for section in report.sections:
        _render_section(doc, section)

    # ── Footer with page numbers ─────────────────────────────────────────
    for sect in doc.sections:
        footer = sect.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = f"{report.title} — {report.date}"
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # ── Save ─────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
